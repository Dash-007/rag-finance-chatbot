import asyncio
import logging
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Dict, List, Any, Optional, Union
from dataclasses import dataclass, field
from datetime import datetime
import hashlib
import mimetypes

# Document processing libraries
import PyPDF2
import pdfplumber
from docx import Document as DocxDocument
import pandas as pd
from bs4 import BeautifulSoup
import markdown

# OCR and image processing
try:
    import pytesseract
    from PIL import Image
    HAS_OCR = True
except ImportError:
    HAS_OCR = False
    
# NLP libraries for content analysis
try:
    import spacy
    HAS_SPACY = True
except ImportError:
    HAS_SPACY: False
    
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class DocumentMetadata:
    """
    Enhanced metadata for enterprise documents.
    """
    file_path: str
    file_name: str
    file_size: int
    file_type: str
    mime_type: str
    document_hash: str
    processed_date: datetime
    source_system: Optional[str] = None
    department: Optional[str] = None
    classification: str = "public" # public, internal, confidential, restricted
    author: Optional[str] = None
    creation_date: Optional[datetime] = None
    language: str = "en"
    document_type: Optional[str] = None # financial_report, legal_contract
    version: str = "1.0"
    tags: List[str] = field(default_factory=list)
    custom_fields: Dict[str, Any] = field(default_factory=dict)
    
@dataclass
class ProcessingResult:
    """
    Result of document processing.
    """
    success: bool
    documents: List[Document]
    metadata: DocumentMetadata
    error: Optional[str] = None
    processing_time: float = 0.0
    extracted_elements: Dict[str, Any] = field(default_factory=dict)
    
class BaseDocumentProcessor(ABC):
    """
    Abstract base class for document processors.
    """
    
    @abstractmethod
    async def process(self, file_path: Path, metadata: Dict[str, Any]) -> ProcessingResult:
        """Process a document and return structured content."""
        pass
    
    @abstractmethod
    def can_process(self, file_path: Path) -> bool:
        """Check if this processor can handle the file type."""
        pass
    
    def _create_metadata(self, file_path: Path, **kwargs) -> DocumentMetadata:
        """Create metadata for a document."""
        file_stat = file_path.stat()
        mime_type, _ = mimetypes.guess_type(str(file_path))
        
        # Generate file hash for deduplication
        with open(file_path, 'rb') as f:
            file_hash = hashlib.sha256(f.read()).hexdigest()
            
        return DocumentMetadata(
            file_path=str(file_path),
            file_name=file_path.name,
            file_size=file_stat.st_size,
            file_type=file_path.suffix.lower(),
            mime_type=mime_type or "unknown",
            document_hash=file_hash,
            processed_date=datetime.now(),
            **kwargs
        )
        
class PDFProcessor(BaseDocumentProcessor):
    """
    Advanced PDF processor with OCR and table extraction.
    """
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            seperators=["\n\n", "\n", ". ", " ", ""]
        )
        
    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".pdf"
    
    async def process(self, file_path: Path, metadata: Dict[str, Any]) -> ProcessingResult:
        """
        Process PDF with advanced extraction capabilities.
        """
        start_time = asyncio.get_event_loop().time()
        
        try:
            # Create metadata
            doc_metadata = self._create_metadata(file_path, **metadata)
            
            # Extract text and structured data
            extracted_content = await self._extract_pdf_content(file_path)
            # Create documents with enhanced metadata
            documents = []
            for i, chunk in enumerate(extracted_content["text_chunks"]):
                if chunk.strip(): # Skip empty chunks
                    doc = Document(
                        page_content=chunk,
                        metadata={
                            **doc_metadata.__dict__,
                            'chunk_id': i,
                            'page_numbers': extracted_content.get('page_mapping', {}).get(i, []),
                            'has_tables': bool(extracted_content.get('tables')),
                            'has_images': bool(extracted_content.get('images'))
                        }
                    )
                    documents.append(doc)
                    
            processing_time = asyncio.get_event_loop().time() - start_time
            
            return ProcessingResult(
                success=True,
                documents=documents,
                metadata=doc_metadata,
                processing_time=processing_time,
                extracted_elements={
                    'tables': extracted_content.get('tables', []),
                    'images': extracted_content.get('image', []),
                    'metdata': extracted_content.get('pdf_metadata', {})
                }
            )
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            processing_time = asyncio.get_event_loop().time() - start_time
            
            return ProcessingResult(
                success=False,
                documents=[],
                metadata=self._create_metadata(file_path, **metadata),
                error=str(e),
                processing_time=processing_time
            )
            
    async def _extract_pdf_content(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract text, tables, and images from PDF.
        """
        content = {
            'text_chunks': [],
            'tables': [],
            'images': [],
            'page_mapping': {},
            'pdf_metadata': {}
        }
        
        try:
            # Use pdfplumber for text and table extraction
            with pdfplumber.open(file_path) as pdf:
                content['pdf_metadata'] = pdf.metadata
                full_text = ""
                
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text
                    page_text = page.extract_text() or ""
                    full_text += f"\n[Page {page_num}]\n" + page_text
                    
                    # Extract table
                    tables = page.extract_tables()
                    for table in tables:
                        if table: # Skip empty tables
                            content['tables'].append({
                                'page': page_num,
                                'data': table,
                                'text_representation': self._table_to_text(table)
                            })
                            
                    # Chunk the text
                    chunks = self.text_splitter.split_text(full_text)
                    content['text_chunks'] = chunks
                    
                    # Create page mapping for chunks
                    content['page_mapping'] = self._create_page_mapping(chunks, full_text)
                    
        except Exception as e:
            # Fallback to PyPDF2 for basic text extraction
            logger.warning(f"pdfplumber failed, falling back to PyPDF2: {e}")
            content = await self._fallback_pdf_extraction(file_path)
            
        return content
    
    def _table_to_text(self, table: List[List[str]]) -> str:
        """
        Convert table data to readable text format.
        """
        if not table:
            return ""
        
        # Find the maximum width for each column
        col_widths = []
        for row in table:
            for i, cell in enumerate(row):
                if i >= len(col_widths):
                    col_widths.append(0)
                if cell:
                    col_widths[i] = max(col_widths[i], len(str(cell)))
                    
        # Format table as text
        formatted_rows = []
        for row in table:
            formatted_row = []
            for i, cell in enumerate(row):
                cell_str = str(cell) if cell else ""
                formatted_row.append(cell_str.ljust(col_widths[i]))
            formatted_rows.append(" | ".join(formatted_row))
            
        return "\n".join(formatted_rows)
    
    def _create_page_mapping(self, chunks: List[str], full_text: str) -> Dict[int, List[int]]:
        """
        Map chunks to page numbers.
        """
        mapping = {}
        # Simple implementation
        for i, chunk in enumerate(chunks):
            # Find page numbers mentioned in chunk
            pages = []
            lines = chunk.split('\n')
            for line in lines:
                if line.startswith('[Page '):
                    try:
                        page_num = int(line.split()[1].strip(']'))
                        pages.append(page_num)
                    except (IndexError, ValueError):
                        continue
            mapping[i] = pages
        return mapping
    
    async def _fallback_pdf_extraction(self, file_path: Path) -> Dict[str, Any]:
        """
        Fallback extraction using PyPDF2.
        """
        content = {'text_chunks': [], 'tables': [], 'images': [], 'page_mapping': {}, 'pdf_metadata': {}}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                content['pdf_metadata'] = pdf_reader.metadata
                
                full_text = ""
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text()
                    full_text += f"\n[Page {page_num}]\n" + page_text
                    
                chunks = self.text_splitter.split_text(full_text)
                content['text_chunks'] = chunks
                content['page_mapping'] = self._create_page_mapping(chunks, full_text)
                
        except Exception as e:
            logger.error(f"Fallback PDF extraction failed: {e}")
            
        return content
    
class WordProcessor(BaseDocumentProcessor):
    """
    Process Microsoft Word documents.
    """
    
    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in ['.docx', '.doc']
    
    async def process(self, file_path: Path, metadata: Dict[str, Any]) -> ProcessingResult:
        """
        Process Word document.
        """
        start_time = asyncio.get_event_loop().time()
        
        try:
            doc_metadata = self._create_metadata(file_path, **metadata)
            
            # Extract content from Word document
            doc = DocxDocument(file_path)
            
            # Extract text
            full_text = ""
            for paragraph in doc.paragraphs:
                full_text += paragraph.text + "\n"
                
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
                
            # Split into chunks
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
            chunks = text_splitter.split_text(full_text)
            
            # Create documents
            documents = []
            for i, chunk in enumerate(chunks):
                if chunk.strip():
                    doc_obj = Document(
                        page_content=chunk,
                        metadata={
                            **doc_metadata.__dict__,
                            'chunk_id': i,
                            "has_tables": bool(tables)
                        }
                    )
                    documents.append(doc_obj)
                    
            processing_time = asyncio.get_event_loop().time() - start_time
            
            return ProcessingResult(
                success=True,
                documents=documents,
                metadata=doc_metadata,
                processing_time=processing_time,
                extracted_elements={'tables': tables}
            )
            
        except Exception as e:
            logger.error(f"Error processing Word document {file_path}: {e}")
            processing_time = asyncio.get_event_loop().time() - start_time
            
            return ProcessingResult(
                success=False,
                documents=[],
                metadata=self._create_metadata(file_path, **metadata),
                error=str(e),
                processing_time=processing_time
            )
            
class ExcelProcessor(BaseDocumentProcessor):
    """
    Process Excel spreadsheets.
    """
    
    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in ['.xlsx', '.xls']
    
    async def process(self, file_path: Path, metadata: Dict[str, Any]) -> ProcessingResult:
        """
        Process Excel file.
        """
        start_time = asyncio.get_event_loop().time()
        
        try:
            doc_metadata = self._create_metadata(file_path, **metadata)
            
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            documents = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Convert DataFrame to text representation
                sheet_text = f"Sheet: {sheet_name}\n"
                sheet_text += df.to_string(index=False)
                
                # Create document for each sheet
                doc = Document(
                    page_content=sheet_text,
                    metadata={
                        **doc_metadata.__dict__,
                        'sheet_name': sheet_name,
                        'rows': len(df),
                        'columns': len(df.columns)
                    }
                )
                documents.append(doc)
                
                processing_time = asyncio.get_event_loop().time() - start_time
                
                return ProcessingResult(
                    success=True,
                    documents=documents,
                    metadata=doc_metadata,
                    processing_time=processing_time,
                    extracted_elements={'sheets': excel_file.sheet_names}
                )
                
        except Exception as e:
            logger.error(f"Error processing Excel file {file_path}: {e}")
            processing_time = asyncio.get_event_loop().time() - start_time
            
            return ProcessingResult(
                success=False,
                documents=[],
                metadata=self._create_metadata(file_path, **metadata),
                error=str(e),
                processing_time=processing_time
            )
            
class EnterpriseDocumentProcessor:
    """
    Main document processing coordinator for enterprise documents.
    """
    
    def __init__(self):
        self.processors = [
            PDFProcessor(),
            WordProcessor(),
            ExcelProcessor()
        ]
        self.processing_stats = {
            'total_processed': 0,
            'successful': 0,
            'failed': 0,
            'by_type': {}
        }
        
    async def process_document(self, file_path: Union[str, Path], metadata: Optional[Dict[str, Any]] = None) -> ProcessingResult:
        """
        Process a single document using the appropriate processor.
        """
        file_path = Path(file_path)
        metadata = metadata or {}
        
        # Find appropriate processor
        processor = self._find_processor(file_path)
        if not processor:
            return ProcessingResult(
                success=False,
                documents=[],
                metadata=DocumentMetadata(
                    file_path=str(file_path),
                    file_name=file_path.name,
                    file_size=0,
                    file_type=file_path.suffix,
                    mime_type="unknown",
                    document_hash="",
                    processed_date=datetime.now()
                ),
                error=f"No processor found for file type: {file_path.suffix}"
            )
            
        # Process document
        result = await processor.process(file_path, metadata)
        
        # Update statistics
        self._update_stats(result)
        
        return result
    
    async def process_directory(self, directory_path: Union[str, Path],
                                metadata: Optional[Dict[str, Any]] = None,
                                recursive: bool = True) -> List[ProcessingResult]:
        """
        Process all documents in a directory.
        """
        directory_path = Path(directory_path)
        results = []
        
        if recursive:
            pattern = "**/*"
        else:
            pattern = "*"
            
        for file_path in directory_path.glob(pattern):
            if file_path.is_file() and self._find_processor(file_path):
                result = await self.process_document(file_path, metadata)
                results.append(result)
                
        return results
        
    def _find_processor(self, file_path: Path) -> Optional[BaseDocumentProcessor]:
        """
        Find the appropriate processor for a file.
        """
        for processor in self.processors:
            if processor.can_process(file_path):
                return processor
        return None
    
    def _update_stats(self, result: ProcessingResult):
        """
        Updating processing statistics.
        """
        self.processing_stats['total_processed'] += 1
        
        if result.success:
            self.processing_stats['successful'] += 1
        else:
            self.processing_stats['failed'] += 1
            
        file_type = result.metadata.file_type
        if file_type not in self.processing_stats['by_type']:
            self.processing_stats['by_type'][file_type] = {'successful': 0, 'failed': 0}
            
        if result.success:
            self.processing_stats['by_type'][file_type]['successful'] += 1
        else:
            self.processing_stats['by_type'][file_type]['failed'] += 1